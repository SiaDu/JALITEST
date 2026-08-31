"""Generate anonymised Study A materials from the production dual-plan pipeline.

This script deliberately calls ``generate_dual_performance_plan`` for every
sample.  It only projects the resulting canonical v2 plan into a read-only
DOCX view; it does not introduce a study-specific LLM prompt or schema.
"""
from __future__ import annotations

import argparse
import json
import os
import re
import shutil
import statistics
import sys
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from zipfile import ZipFile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
sys.path[:0] = [str(ROOT / "src"), str(ROOT / "tools" / "maya")]

from expregaze_jali.transcript_anchor_model import build_conversation_anchor_model  # noqa: E402
from dual_sparse_score_model import build_dialogue_projection, render_actor_score, render_initial_score  # noqa: E402

DEFAULT_PROMPT_TEMPLATE = ROOT / "prompts" / "actor_dual_semantic_beat_prompt_v1.md"

SCENES = {
    "S1": ("SOL", "DANA", "Dana and Sol are colleagues working on a prototype. The prototype has failed several times during the week, but Sol has now repaired the problem. There is no unresolved disagreement between them.", "SOL: I thought it was going to fail again.\nDANA: It didn't.\nSOL: You think it's ready?\nDANA: Yes. You did a good job."),
    "S2": ("ERIN", "KAI", "Erin and Kai are collaborating on the same project. Erin had explicitly asked Kai not to send their unfinished draft before they reviewed it together. Kai sent it anyway because the deadline was approaching. They still need to continue working together.", "ERIN: You sent the draft.\nKAI: We were running out of time.\nERIN: You should have asked me.\nKAI: I know.\nERIN: Don't do it again."),
    "S3": ("LEO", "MARA", "Earlier that evening, Mara secretly allowed Evan to hide in a locked storage room. Leo is now searching the building for Evan. Mara knows exactly where Evan is but does not want Leo to find him.", "LEO: We're looking for Evan. Have you seen him tonight?\nMARA: No.\nLEO: You're sure?\nMARA: Yes."),
    "S4": ("BEN", "ARI", "Ari and Ben are close colleagues. Both applied for the same lead position. Ari had expected to receive it, but Ben was selected instead. They will continue working together on the same team.", "BEN: They told me this morning. I got the lead.\nARI: I heard.\nBEN: Are we okay?\nARI: Of course. Congratulations."),
    "S5": ("ELI", "NORA", "Nora and Eli met for the first time earlier that evening. Their original conversation has already ended, but both have found unnecessary reasons to continue talking rather than leave. Neither has directly acknowledged any personal interest in the other.", "ELI: I should probably go.\nNORA: Probably.\nELI: Unless you need help with those boxes.\nNORA: I can manage.\nELI: Right.\nNORA: You can stay a minute, though."),
    "S6": ("JULES", "ROWAN", "Rowan and Jules were once very close but stopped speaking after a serious argument several months ago. This is their first calm conversation since the disagreement. They have discussed what happened, but neither has apologized and the conflict is not fully resolved.", "JULES: I should go.\nROWAN: Okay.\nJULES: Maybe I'll come by next week.\nROWAN: I'd like that."),
}

# First item of each pair is the anonymous Plan A source.  The four base rows
# are duplicated exactly as specified for U09--U16.
_ROWS = [
    [("S1D1","S1P1"),("S2P2","S2D2"),("S3D3","S3P3"),("S4P4","S4D4"),("S5D1","S5P1"),("S6P2","S6D2")],
    [("S1D2","S1P2"),("S2P3","S2D3"),("S3D4","S3P4"),("S4P1","S4D1"),("S5D2","S5P2"),("S6P3","S6D3")],
    [("S1D3","S1P3"),("S2P4","S2D4"),("S3D1","S3P1"),("S4P2","S4D2"),("S5D3","S5P3"),("S6P4","S6D4")],
    [("S1D4","S1P4"),("S2P1","S2D1"),("S3D2","S3P2"),("S4P3","S4D3"),("S5D4","S5P4"),("S6P1","S6D1")],
    [("S1P1","S1D1"),("S2D2","S2P2"),("S3P3","S3D3"),("S4D4","S4P4"),("S5P1","S5D1"),("S6D2","S6P2")],
    [("S1P2","S1D2"),("S2D3","S2P3"),("S3P4","S3D4"),("S4D1","S4P1"),("S5P2","S5D2"),("S6D3","S6P3")],
    [("S1P3","S1D3"),("S2D4","S2P4"),("S3P1","S3D1"),("S4D2","S4P2"),("S5P3","S5D3"),("S6D4","S6P4")],
    [("S1P4","S1D4"),("S2D1","S2P1"),("S3P2","S3D2"),("S4D3","S4P3"),("S5P4","S5D4"),("S6D1","S6P1")],
]
ASSIGNMENTS = {f"U{i + 1:02d}": _ROWS[i % 8] for i in range(16)}

def _utc() -> str: return datetime.now(timezone.utc).isoformat()

def _shade(cell: Any, fill: str) -> None:
    props = cell._tc.get_or_add_tcPr(); shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), fill); props.append(shade)

def _keep_row_together(row: Any) -> None:
    props = row._tr.get_or_add_trPr(); props.append(OxmlElement("w:cantSplit"))

def _configure_document(doc: Document) -> None:
    section = doc.sections[0]; section.page_width = Inches(8.27); section.page_height = Inches(11.69)
    section.top_margin = section.bottom_margin = Inches(0.58); section.left_margin = section.right_margin = Inches(0.6)
    normal = doc.styles["Normal"]; normal.font.name = "Arial"; normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(2); normal.paragraph_format.line_spacing = 1.0

def _paragraph(container: Any, text: str = "", *, bold: bool = False, size: float = 9.5,
               color: str | None = None, code: bool = False, keep: bool = False) -> Any:
    p = container.add_paragraph(); p.paragraph_format.space_after = Pt(2); p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.keep_with_next = keep
    run = p.add_run(text); run.bold = bold; run.font.size = Pt(size); run.font.name = "Consolas" if code else "Arial"
    if color: run.font.color.rgb = RGBColor.from_string(color)
    return p

def _score_paragraph(container: Any, runs: list[tuple[str, dict[str, Any]]]) -> None:
    p = container.add_paragraph(); p.paragraph_format.space_after = Pt(3); p.paragraph_format.line_spacing = 1.0
    for value, style in runs:
        run = p.add_run(value); run.bold = bool(style.get("bold")); run.font.name = "Consolas"; run.font.size = Pt(9)
        if style.get("color"): run.font.color.rgb = RGBColor.from_string(style["color"])

def _clear_initial_paragraph(cell: Any) -> None:
    cell.paragraphs[0]._element.getparent().remove(cell.paragraphs[0]._element)

def _score_runs(score: str, projection: Any, actor: str) -> list[tuple[str, dict[str, Any]]]:
    """Apply the same tag and speaking/listening semantics as Maya's highlighter."""
    out: list[tuple[str, dict[str, Any]]] = []; plain = 0
    for part in re.split(r"(<[^<>\r\n]+>)", score):
        if not part: continue
        if part.startswith("<"): out.append((part, {"mono": True, "bold": True, "color": "C026D3"})); continue
        start = 0; active_speaker: str | None = None
        for pos, char in enumerate(part):
            speaker = next((x.speaker for x in projection.speaker_ranges if x.start <= plain + pos < x.end), None)
            if pos and speaker != active_speaker:
                color = "B58900" if str(active_speaker).casefold() == actor.casefold() else "2563EB"
                out.append((part[start:pos], {"mono": True, "color": color})); start = pos
            active_speaker = speaker
        if part:
            color = "B58900" if str(active_speaker).casefold() == actor.casefold() else "2563EB"
            out.append((part[start:], {"mono": True, "color": color}))
        plain += len(part)
    return out

def _anchor_phrase(model: Any, event: dict[str, Any]) -> str:
    anchor = {a.anchor_id: a for turn in model.turns for a in turn.anchors}[event["anchor_id"]]
    turn = next(t for t in model.turns if any(a.anchor_id == anchor.anchor_id for a in t.anchors))
    return turn.utterance_text

def add_plan(doc: Document, plan: dict[str, Any], script: str, actors: tuple[str, str]) -> None:
    model = build_conversation_anchor_model(script, character_a=actors[0], character_b=actors[1]); projection = build_dialogue_projection(model)
    card = doc.add_table(rows=1, cols=1); card.autofit = False; _keep_row_together(card.rows[0])
    cell = card.cell(0, 0); _clear_initial_paragraph(cell); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    _shade(cell, "FFFFFF")
    header = cell.add_table(rows=1, cols=1); hcell = header.cell(0, 0); _shade(hcell, "262626"); _clear_initial_paragraph(hcell)
    _paragraph(hcell, "SEMANTIC PERFORMANCE TAG", bold=True, size=9, color="FFFFFF")
    for actor in plan["characters"]:
        block = cell.add_table(rows=1, cols=1); bcell = block.cell(0, 0); _clear_initial_paragraph(bcell)
        bcell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        actor_head = bcell.add_table(rows=1, cols=1); ah = actor_head.cell(0, 0); _shade(ah, "3A3A3A"); _clear_initial_paragraph(ah)
        _paragraph(ah, f"{actor} PERFORMANCE", bold=True, size=10.5, color="FFFFFF")
        _paragraph(bcell, "INITIAL PERFORMANCE", bold=True, size=8.5, color="666666", keep=True)
        initial = bcell.add_table(rows=1, cols=1).cell(0, 0); _shade(initial, "1F2937"); _clear_initial_paragraph(initial)
        _paragraph(initial, render_initial_score(plan, actor) or "<No initial tags>", bold=True, size=9, color="C026D3", code=True)
        _paragraph(bcell, "DIALOGUE PERFORMANCE", bold=True, size=8.5, color="666666", keep=True)
        score = bcell.add_table(rows=1, cols=1).cell(0, 0); _shade(score, "1F2937"); _clear_initial_paragraph(score)
        _score_paragraph(score, _score_runs(render_actor_score(plan, projection, actor), projection, actor))
        _paragraph(bcell, "REASON BY PHRASE", bold=True, size=8.5, color="666666", keep=True)
        reasons: list[tuple[str, str]] = [("Initial", str(plan["initial_reasons"][actor]))]
        seen = set(reasons)
        for event in plan["tracks"][actor]:
            item = (f'"{_anchor_phrase(model, event)}"', str(event["reason"]))
            if item not in seen: reasons.append(item); seen.add(item)
        table = bcell.add_table(rows=1, cols=2); table.style = "Table Grid"; table.columns[0].width = Inches(1.75); table.columns[1].width = Inches(3.25)
        for index, label in enumerate(("Phrase", "Acting Interpretation")):
            tc = table.cell(0, index); _shade(tc, "E5E7EB"); _clear_initial_paragraph(tc); _paragraph(tc, label, bold=True, size=8.5)
        for phrase, reason in reasons:
            cells = table.add_row().cells
            for tc, value, is_phrase in ((cells[0], phrase, True), (cells[1], reason, False)):
                _clear_initial_paragraph(tc); _paragraph(tc, value, bold=is_phrase, size=8.8)

def mock_plan(scene: str) -> dict[str, Any]:
    a, b, _context, script = SCENES[scene]; model = build_conversation_anchor_model(script, character_a=a, character_b=b)
    anchors = [x for t in model.turns for x in t.anchors]
    return {"schema_version": "dual_performance_plan_v2", "characters": [a, b],
        "initial_states": {a: {"affect": "Neutral-55", "gaze": f"GAZE-{b}"}, b: {"affect": "Neutral-55", "gaze": f"GAZE-{a}"}},
        "initial_reasons": {a: "Holding a clear, responsive presence.", b: "Staying attentive to the exchange."},
        "tracks": {a: [{"anchor_id": anchors[0].anchor_id, "changes": {"affect": "Neutral-60"}, "reason": "Responding carefully to the opening."}], b: [{"anchor_id": anchors[-1].anchor_id, "changes": {"affect": "Neutral-60"}, "reason": "Keeping the exchange measured."}]}}

def _validate_assignments(replicates: int) -> None:
    if replicates != 4: raise ValueError("The supplied counterbalanced assignment is defined for exactly four replicates.")
    for scene in SCENES:
        pair_counts = {i: 0 for i in range(1, 5)}; d_first = 0
        for rows in ASSIGNMENTS.values():
            first, second = next(pair for pair in rows if pair[0][:2] == scene)
            if first[3] != second[3]: raise AssertionError(f"Mismatched replicate pair: {first}, {second}")
            pair_counts[int(first[3])] += 1; d_first += int(first[2] == "D")
        assert pair_counts == {1: 4, 2: 4, 3: 4, 4: 4} and d_first == 8

def _forbidden(text: str) -> bool:
    return bool(re.search(r"SEMANTIC BEATS|\[BEATS\]|\bw\d{4}\b|S[1-6][DP][1-4]|Dialogue-Only|Full-Context|replicate|gpt-|run_id", text, re.I))

def _read_docx_text(path: Path) -> str:
    with ZipFile(path) as archive:
        return re.sub(r"<[^>]+>", " ", archive.read("word/document.xml").decode("utf-8"))

def _explicit_page_break_count(path: Path) -> int:
    with ZipFile(path) as archive:
        return archive.read("word/document.xml").count(b'w:type="page"')

def _prepare_output(path: Path, overwrite: bool) -> None:
    if path.exists() and any(path.iterdir()):
        if not overwrite: raise FileExistsError(f"Refusing to overwrite existing StudyA directory: {path}. Use --overwrite.")
        backup = path.with_name(path.name + ".backup_" + datetime.now().strftime("%Y%m%d_%H%M%S")); shutil.move(str(path), str(backup))
        print(f"Backed up existing output to: {backup}")
    path.mkdir(parents=True, exist_ok=True)

def _render_participants(out: Path, plans: dict[str, dict[str, Any]], *, users: tuple[str, ...] | None = None,
                         target_dir: Path | None = None) -> None:
    for user in users or tuple(ASSIGNMENTS):
        rows = ASSIGNMENTS[user]
        doc = Document(); _configure_document(doc)
        for index, (left, right) in enumerate(rows, 1):
            scene = f"S{index}"; a, b, context, script = SCENES[scene]
            _paragraph(doc, f"SCENE {index}", bold=True, size=13, keep=True)
            _paragraph(doc, "Context", bold=True, size=10.5, keep=True)
            context_box = doc.add_table(rows=1, cols=1).cell(0, 0); _shade(context_box, "F1F3F5"); _clear_initial_paragraph(context_box); _paragraph(context_box, context, size=9)
            _paragraph(doc, "Dialogue", bold=True, size=10.5, keep=True)
            dialogue_box = doc.add_table(rows=1, cols=1).cell(0, 0); _clear_initial_paragraph(dialogue_box)
            for line in script.splitlines():
                speaker, utterance = line.split(": ", 1); para = _paragraph(dialogue_box, "", size=9); para.add_run(speaker + ": ").bold = True; para.add_run(utterance)
            for label, source in (("PLAN A", left), ("PLAN B", right)):
                if label == "PLAN B":
                    _paragraph(doc, f"SCENE {index}", bold=True, size=11, keep=True)
                    _paragraph(doc, "Context", bold=True, size=9.5, keep=True)
                    context_box = doc.add_table(rows=1, cols=1).cell(0, 0); _shade(context_box, "F1F3F5"); _clear_initial_paragraph(context_box); _paragraph(context_box, context, size=8.5)
                    _paragraph(doc, "Dialogue", bold=True, size=9.5, keep=True)
                    dialogue_box = doc.add_table(rows=1, cols=1).cell(0, 0); _clear_initial_paragraph(dialogue_box)
                    for line in script.splitlines():
                        speaker, utterance = line.split(": ", 1); para = _paragraph(dialogue_box, "", size=8.5); para.add_run(speaker + ": ").bold = True; para.add_run(utterance)
                _paragraph(doc, label, bold=True, size=15.5, keep=True)
                add_plan(doc, plans[source], script, (a, b))
                if label == "PLAN A" or index < 6: doc.add_page_break()
        target = (target_dir or out / "participants") / f"{user}.docx"; target.parent.mkdir(parents=True, exist_ok=True); doc.save(target)
        text = _read_docx_text(target)
        assert not _forbidden(text), f"Participant leak in {target}"
        assert text.count("REASON BY PHRASE") == 24
        assert text.count("INITIAL PERFORMANCE") == 24 and text.count("DIALOGUE PERFORMANCE") == 24

def _timing_stats(samples: list[dict[str, Any]]) -> dict[str, float]:
    values = [x["llm_runtime_sec"] for x in samples]
    return {"total_llm_runtime_sec": sum(values), "mean_llm_runtime_sec": statistics.mean(values), "median_llm_runtime_sec": statistics.median(values), "min_llm_runtime_sec": min(values), "max_llm_runtime_sec": max(values), "stdev_llm_runtime_sec": statistics.stdev(values) if len(values) > 1 else 0.0}

def _load_dotenv_for_diagnostics(dotenv_path: Path) -> None:
    """Load the repository .env without revealing any value in diagnostics."""
    try:
        from dotenv import load_dotenv
        load_dotenv(dotenv_path=dotenv_path)
        return
    except ImportError:
        # Keep --dry-run dependency-light while following dotenv's ordinary
        # KEY=VALUE convention.  The production runner remains authoritative
        # for real API validation and loading.
        pass
    if not dotenv_path.exists():
        return
    for line in dotenv_path.read_text(encoding="utf-8").splitlines():
        match = re.match(r"\s*(?:export\s+)?([A-Za-z_][A-Za-z0-9_]*)\s*=\s*(.*)\s*$", line)
        if not match or line.lstrip().startswith("#"):
            continue
        value = match.group(2).strip().strip('"').strip("'")
        os.environ.setdefault(match.group(1), value)

def _configured_api_key_env(config_path: Path) -> str:
    """Read this one scalar without importing YAML for a mock-only export."""
    match = re.search(r"^\s*api_key_env\s*:\s*([^#\s]+)", config_path.read_text(encoding="utf-8"), re.MULTILINE)
    return match.group(1).strip("'\"") if match else "OPENAI_API_KEY"

def _print_diagnostics(config_path: Path) -> tuple[Path, str, bool]:
    dotenv_path = ROOT / ".env"
    _load_dotenv_for_diagnostics(dotenv_path)
    api_key_env = _configured_api_key_env(config_path)
    available = bool(os.getenv(api_key_env))
    print(f"repo root: {ROOT}")
    print(f"Python executable: {sys.executable}")
    print(f"cwd: {Path.cwd()}")
    print(f".env file exists: {dotenv_path.exists()}")
    print(f"configured api_key_env: {api_key_env}")
    print(f"API key available after dotenv load: {available}")
    return dotenv_path, api_key_env, available

def main() -> None:
    p = argparse.ArgumentParser(); p.add_argument("--output-dir", type=Path, default=Path.home() / "Desktop" / "StudyA"); p.add_argument("--replicates", type=int, default=4); p.add_argument("--overwrite", action="store_true"); p.add_argument("--dry-run", action="store_true"); p.add_argument("--llm-config", type=Path, default=ROOT / "configs" / "llm.yaml"); args = p.parse_args()
    _validate_assignments(args.replicates); _prepare_output(args.output_dir, args.overwrite)
    _print_diagnostics(args.llm_config)
    if args.dry_run:
        plans = {f"{scene}{condition}{rep}": mock_plan(scene) for scene in SCENES for condition in "DP" for rep in range(1, 5)}
        preview_path = args.output_dir / "U01.docx"
        _render_participants(args.output_dir, plans, users=("U01",), target_dir=args.output_dir)
        assert preview_path.exists(), f"Preview DOCX was not created: {preview_path}"
        assert _explicit_page_break_count(preview_path) == 11, "Preview must deliberately create 12 pages."
        (args.output_dir / "README.txt").write_text("DRY RUN ONLY: participant DOCX files use mock plans; no API calls were made.\n", encoding="utf-8")
        print(f"preview DOCX absolute path: {preview_path.resolve()}")
        print(f"preview DOCX exists: {preview_path.exists()}")
        print(f"Mock preview DOCX generated without API calls: {preview_path}"); return
    plans: dict[str, dict[str, Any]] = {}; samples: list[dict[str, Any]] = []; failures: list[dict[str, Any]] = []
    for scene, (a, b, context, script) in SCENES.items():
        for condition in "DP":
            for rep in range(1, 5):
                code = f"{scene}{condition}{rep}"; run_dir = args.output_dir / "raw" / code; attempt = 0; sample_failures: list[dict[str, Any]] = []
                while True:
                    attempt += 1; started = _utc(); wall = time.perf_counter(); call_times: list[float] = []
                    # Import the production generator only for real calls.  This
                    # keeps --dry-run usable in a minimal Python environment.
                    from expregaze_jali.generate_dual_performance_plan import generate_dual_performance_plan
                    from expregaze_jali.generate_performance_plan import hci_run_paths
                    from expregaze_jali.run_actor_llm import generate_text_artifacts
                    def timed_runner(**kwargs: Any) -> tuple[str, dict[str, Any]]:
                        tic = time.perf_counter()
                        try: return generate_text_artifacts(**kwargs)
                        finally: call_times.append(time.perf_counter() - tic)
                    try:
                        plan = generate_dual_performance_plan(script=script, character_a=a, character_b=b, context=context if condition == "P" else None, run_id=code.lower(), output_dir=run_dir, prompt_template_path=DEFAULT_PROMPT_TEMPLATE, llm_config_path=args.llm_config, overwrite=True, proposal_runner=timed_runner)
                        elapsed = call_times[-1]
                        response_meta_path = hci_run_paths(code.lower(), run_dir).response_meta
                        response_meta = json.loads(response_meta_path.read_text(encoding="utf-8"))
                        item = {"scene_id": scene, "condition": condition, "replicate": rep, "run_id": code.lower(), "model_name": response_meta.get("model"), "model_config_path": str(args.llm_config), "model_config": args.llm_config.read_text(encoding="utf-8"), "prompt_template_path": str(DEFAULT_PROMPT_TEMPLATE), "generation_started_at": started, "generation_finished_at": _utc(), "llm_runtime_sec": elapsed, "total_run_runtime_sec": time.perf_counter() - wall, "attempt_count": attempt, "failed_attempts": sample_failures}
                        (run_dir / "participant_score_data.json").write_text(json.dumps({"plan": plan}, indent=2), encoding="utf-8")
                        plans[code] = plan; samples.append(item); break
                    except Exception as exc:
                        failed = {"source": code, "attempt": attempt, "llm_runtime_sec": call_times[-1] if call_times else 0.0, "total_run_runtime_sec": time.perf_counter() - wall, "error": str(exc), "validation_status": "failed"}
                        failures.append(failed); sample_failures.append(failed)
                        if attempt >= 3: raise RuntimeError(f"{code} failed after three attempts") from exc
                        if run_dir.exists():
                            failed_dir = args.output_dir / "raw" / "failed_attempts" / f"{code}_attempt{attempt}"
                            failed_dir.parent.mkdir(parents=True, exist_ok=True)
                            shutil.move(str(run_dir), str(failed_dir))
    _render_participants(args.output_dir, plans); stats = _timing_stats(samples)
    manifest = {"successful_generations": len(samples), "samples": samples, "failed_attempts": failures, "timing_statistics": stats}
    (args.output_dir / "generation_manifest.json").write_text(json.dumps(manifest, indent=2) + "\n", encoding="utf-8")
    (args.output_dir / "assignment_manifest.json").write_text(json.dumps(ASSIGNMENTS, indent=2) + "\n", encoding="utf-8")
    (args.output_dir / "README.txt").write_text("Study A materials generated from the production dual Performance Plan pipeline.\n" + json.dumps(stats, indent=2) + "\n", encoding="utf-8")
    assert len(plans) == 48 and len(list((args.output_dir / "participants").glob("U*.docx"))) == 16
    print(f"Generated {len(plans)} plans in {args.output_dir}")

if __name__ == "__main__": main()
